import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
import os
from io import BytesIO
from dashboard import load_sheet
def get_val(row, col):
    # direct match
    if col in row:
        return row[col]

    # possible variations
    variants = [
        col.replace("-", "_"),
        col.replace("_", "-"),
        col.replace("-", "_").replace("__", "_"),
        col.replace("-", "_").replace("_", "__"),
        col.replace("-", "_").replace("_", "_"),
        col.replace("agri_practices-", "agri_practices_-"),   # 🔥 KEY FIX
        col.replace("agri_practices_", "agri_practices_-"),   # 🔥 KEY FIX
    ]

    for c in variants:
        if c in row:
            return row[c]

    return 0

st.title("📄 Village Reports")

df_profile = load_sheet("village profile")
df_plan = load_sheet("village plan")
df_budget = load_sheet("budget")

village_list = sorted(df_profile["village"].dropna().unique())
selected_village = st.selectbox("Select Village", village_list)

df_v = df_profile[df_profile["village"] == selected_village]
df_p = df_plan[df_plan["village"] == selected_village]


# ---------------- REPORT TEXT ----------------
def generate_report(df_v, df_p, village):

    text = "Required Works:\n"

    for col in df_p.columns:
        if "required" in col.lower():
            val = df_p[col].sum()

            if val > 0:
                clean_name = col.replace("_", " ").replace("required", "").replace("farmeasy", "").strip().title()
                text += f"\n- {clean_name}: {int(val)} units required"

    return text

# ---------------- DOC CREATION ----------------
def create_doc(text, df_v, village):

    doc = Document()
    row = df_v.iloc[0]

    # TITLE
    doc.add_heading("Village Report", 0)

    # LOCATION
    location = f"{row.get('village_gps-Latitude','')}, {row.get('village_gps-Longitude','')}"

    # TABLE
    data = [
        ("Village", village),
        ("Location (Lat, Long)", location),
        ("Mandal", row.get("mandal", "")),
        ("Panchayath", row.get("panchayath", "")),
        ("VO Name", row.get("VO_name", "")),
        ("Raithu Seva Kendra", row.get("RSK_name", "")),
        ("Sachivalayam", row.get("RSK_name", ""))
    ]

    table = doc.add_table(rows=len(data), cols=2)
    table.style = "Table Grid"

    for i, (key, val) in enumerate(data):
        table.rows[i].cells[0].text = str(key)
        table.rows[i].cells[1].text = str(val)

    # SPACE
    doc.add_paragraph("")

    # -------- VALUES --------
    total_pop = row.get("population", "NA")
    male = row.get("Households-male", "NA")
    female = row.get("Households-femlae", "NA")
    hh = row.get("Total HHs", "NA")

    children_icds = row.get("children_icds", "NA")
    children_school = row.get("children_school", "NA")

    school = str(row.get("school", "")).strip().lower()
    school_name = row.get("school_name", "")
    kg = str(row.get("kg_school", "")).strip().lower()

    water = row.get("drinking_water_source", "NA")

    # -------- JOB CARDS --------

    def clean_int(val):
        try:
            return int(float(str(val).strip()))
        except:
            return 0

    total_hhs = clean_int(row.get("Total HHs"))
    job_yes = clean_int(row.get("Households-mnregs_cards"))

    # 👉 ALWAYS derive job_no
    job_no = max(total_hhs - job_yes, 0)
    st.write("DEBUG JOB:", 
         row.get("Total HHs"), 
         row.get("Households-mnregs_cards"), 
         row.get("No of HHs not having Job cards"))

    # -------- SMART SENTENCE --------
    if total_hhs > 0 and job_yes == total_hhs:
        job_text = f"All {total_hhs} households in the village possess job cards, and no households are left out."

    elif job_yes > 0 and job_yes < total_hhs:
        job_text = f"In terms of livelihoods, {job_yes} households possess job cards, while {job_no} households do not have access to them."

    elif job_yes == 0 and total_hhs > 0:
        job_text = "No households in the village currently possess job cards, indicating a need for increased access."

    else:
        job_text = "Job card information is not available."
    
    # MIGRATION
    mig_hhs = pd.to_numeric(row.get("HHs going for seasonal migraion", 0), errors="coerce")
    mig_members = pd.to_numeric(row.get("Households-migration_members", 0), errors="coerce")
    mig_days = pd.to_numeric(row.get("Average no of days in a year going for migraion", 0), errors="coerce")
    mig_income = pd.to_numeric(row.get("Average earning per annum per family", 0), errors="coerce")

    mig_hhs = int(mig_hhs) if pd.notna(mig_hhs) else 0
    mig_members = int(mig_members) if pd.notna(mig_members) else 0
    mig_days = int(mig_days) if pd.notna(mig_days) else 0
    mig_income = int(mig_income) if pd.notna(mig_income) else 0
    
    # -------- MIGRATION TEXT --------
    if mig_hhs == 0 or mig_members == 0:
        migration_text = "No households from the village undertake seasonal migration."
    else:
        migration_text = f"""{mig_hhs} households undertake seasonal migration involving {mig_members} members. On average, migration lasts for about {mig_days} days per year, with an average annual earning of ₹{mig_income} per family."""

    # -------- PARAGRAPH --------
    if school in ["yes", "y"]:
        para = f"""{village} is a small village with a total population of {total_pop} people, comprising {male} males and {female} females across {hh} households. The village has {children_icds} children enrolled in ICDS and {children_school} children attending {school_name}. A kitchen garden is {'available' if kg in ['yes','y'] else 'not available'} at the school. Drinking water in the village is sourced from {water}. In terms of livelihoods, {job_yes} households possess job cards, while {job_no} households do not have access to them.

    {migration_text}
    """
    else:
        para = f"""{village} is a small village with a total population of {total_pop} people, comprising {male} males and {female} females across {hh} households. The village has {children_icds} children enrolled in ICDS; however, there is no functioning school currently, and no children are attending school despite the presence of a {school_name}. Drinking water in the village is sourced from {water}. In terms of livelihoods, {job_yes} households possess job cards, while {job_no} households do not have access to them.

    {migration_text}
    """

    doc.add_paragraph(para)

    # -------- COMMUNITY TABLE --------
    doc.add_heading("Community-wise Households", 1)

    others_name = row.get("Households-community_others", "Others")

    raw_data = [
        ("Kotia", row.get("Households-Kotia", 0)),
        ("Porja", row.get("Households-Porja", 0)),
        ("Kondadora", row.get("Households-Kondadora", 0)),
        ("Nookadora", row.get("Households-Nookadora", 0)),
        ("Kammari", row.get("Households-Kammari", 0)),
        ("Bhagatha", row.get("Households-Bhagatha", 0)),
        ("Valmiki", row.get("Households-Valmiki", 0)),
        ("Kondhu PVTG", row.get("Households-Kondu_PVTG", 0)),
        (others_name, row.get("Households-community_others_hhs", 0)),
    ]

    clean_data = []
    total = 0

    for name, val in raw_data:
        val = pd.to_numeric(val, errors="coerce")
        val = int(val) if pd.notna(val) else 0

        if val > 0:   # ✅ skip zero rows
            clean_data.append((name, val))
            total += val

    # always add total
    clean_data.append(("Total", total))

    # create table
    table2 = doc.add_table(rows=len(clean_data), cols=2)
    table2.style = "Table Grid"

    for i, (name, val) in enumerate(clean_data):
        table2.rows[i].cells[0].text = str(name)
        table2.rows[i].cells[1].text = str(val)

# -------- AGRICULTURE: LAND DETAILS --------
    doc.add_heading("Agriculture - Land Details", 1)

    land_data_raw = [
        ("Mettu Land (acres)", row.get("mettu_total_land_acr", 0)),
        ("Pallam Land (acres)", row.get("Total pallam land", 0)),
        ("Podu Land (acres)", row.get("podu_total_land_acr", 0)),
        ("Banjaru Land (acres)", row.get("banjaru-acres", 0)),
        ("Coffee/Cashew Land (acres)", row.get("coffee_cashew_land_acr", 0)),
        ("Total Land (acres)", row.get("Total land in the village_acre", 0)),
    ]

    land_data = []

    for name, val in land_data_raw:
        val = pd.to_numeric(val, errors="coerce")
        val = int(val) if pd.notna(val) else 0

        if val > 0:
            land_data.append((name, val))

    # create table only if data exists
    if land_data:
        table3 = doc.add_table(rows=len(land_data), cols=2)
        table3.style = "Table Grid"

        for i, (name, val) in enumerate(land_data):
            table3.rows[i].cells[0].text = name
            table3.rows[i].cells[1].text = str(val)
    else:
        doc.add_paragraph("No land data available.")

    # -------- AGRICULTURE PARAGRAPH --------
    doc.add_paragraph("")

    farming_hhs = pd.to_numeric(row.get("Households-farming_hhs", 0), errors="coerce")
    landless = pd.to_numeric(row.get("Total no of land less HHs", 0), errors="coerce")

    kharif_farmers = pd.to_numeric(row.get("Crops-kharif_farmers", 0), errors="coerce")
    kharif_acres = pd.to_numeric(row.get("Crops-kharif_acres", 0), errors="coerce")

    rabi_farmers = pd.to_numeric(row.get("Crops-rabi_farmers", 0), errors="coerce")
    rabi_acres = pd.to_numeric(row.get("Crops-rabi_acres", 0), errors="coerce")

    irrig_sources = row.get("Crops-irrigation_sources", "NA")
    kharif_crops = row.get("Crops-crops_kharif", "NA")
    rabi_crops = row.get("Crops-crops_rabi", "NA")

    rainfed = pd.to_numeric(row.get("Crops-rainfed_acrs", 0), errors="coerce")
    irrigated = pd.to_numeric(row.get("Crops-irrigated_acrs", 0), errors="coerce")

    new_irrig = row.get("Crops-new_irrigation_source", "NA")
    irrig_type = row.get("Crops-potential_irrigation", "NA")

    grazing = row.get("Crops-Grazing", "NA")
    ntfp = row.get("Crops-NTFP", "NA")

# safe int conversion
    def safe_int(val):
        return int(val) if pd.notna(val) else 0

    farming_hhs = safe_int(farming_hhs)
    landless = safe_int(landless)
    kharif_farmers = safe_int(kharif_farmers)
    kharif_acres = safe_int(kharif_acres)
    rabi_farmers = safe_int(rabi_farmers)
    rabi_acres = safe_int(rabi_acres)
    rainfed = safe_int(rainfed)
    irrigated = safe_int(irrigated)

    agri_para = f"""Agriculture is the primary livelihood activity in the village, with {farming_hhs} households engaged in farming, while {landless} households are landless. 

    During the Kharif season, {kharif_farmers} farmers cultivate crops over an extent of {kharif_acres} acres, mainly growing {kharif_crops}. In the Rabi season, {rabi_farmers} farmers cultivate crops over {rabi_acres} acres, with crops such as {rabi_crops}. 

    The village largely depends on rainfed agriculture covering {rainfed} acres, with only {irrigated} acres under irrigation. Major irrigation sources include {irrig_sources}. There is {'availability' if str(new_irrig).lower()=='yes' else 'no availability'} of new irrigation sources, and potential irrigation options include {irrig_type}. 

    Livestock grazing follows a {grazing} system, and Non-Timber Forest Products (NTFP) such as {ntfp} are available in the village, contributing to livelihoods."""

    doc.add_paragraph(agri_para)

    # -------- CROP MODELS TABLE (WITH YIELD) --------
    doc.add_heading("Crop Models Practiced in the Village", 1)

    crop_models_raw = [
        ("Cashew Mono",
         get_val(row, "agri_practices-Cashew_mono_acres"),
         get_val(row, "agri_practices-Cashew_mono_yield_qntl_acr")),

        ("Cashew Poly",
         get_val(row, "agri_practices-Cashew_poly_acres"),
         get_val(row, "agri_practices-Cashew_poly_yield_qntl_acr")),

        ("Mango Mono",
         get_val(row, "agri_practices-Mango_mono_acres"),
         get_val(row, "agri_practices-Mango_mono_yield_qntl_acr")),

        ("Mango Poly",
         get_val(row, "agri_practices-Mango_poly_acres"),
         get_val(row, "agri_practices-Mango_poly_yield_qntl_acr")),

        ("Coffee Mono",
         get_val(row, "agri_practices-Coffee_mono_acres"),
         get_val(row, "agri_practices-Coffee_mono_yield_qntl_acr")),

        ("Coffee with Pepper",
         get_val(row, "agri_practices-Coffee_with_pepper_acres"),
         get_val(row, "agri_practices-Coffee_with_pepper_yield_qntl_acr")),

        ("Millet Broadcasting",
         get_val(row, "agri_practices-Millet_broadcasting_acres"),
         get_val(row, "agri_practices-Millet_broadcasting_yield_qntl_acr")),

        ("Millet Line Sowing",
         get_val(row, "agri_practices-Millet_linesowing_acres"),
         get_val(row, "agri_practices-Millet_linesowing_yield_qntl_acr")),

        ("Guliragi Mono",
         get_val(row, "agri_practices-Guliragi_acres"),
         get_val(row, "agri_practices-Guliragi_yield_qntl_acr")),

        ("Guliragi Poly",
         get_val(row, "agri_practices-Guliragi_poly_acres"),
         get_val(row, "agri_practices-Guliragi_poly_yield_qntl_acr")),

        ("Sirisama Mono",
         get_val(row, "agri_practices-Sirisama_mono_acres"),
         get_val(row, "agri_practices-Sirisama_mono_yield_qntl_acr")),

        ("Sirisama Poly",
         get_val(row, "agri_practices-Sirisama_poly_acres"),
         get_val(row, "agri_practices-Sirisama_poly_yield_qntl_acr")),

        ("SRI Paddy",
         get_val(row, "agri_practices-SRI_paddy_acres"),
         get_val(row, "agri_practices-SRI_paddy_yield_qntl_acr")),

        ("Paddy Line Sowing",
         get_val(row, "agri_practices-Paddy_linesowing_acres"),
         get_val(row, "agri_practices-Paddy_linesowing_yield_qntl_acr")),

        ("Ginger Mono",
         get_val(row, "agri_practices-Ginger_mono_acres"),
         get_val(row, "agri_practices-Ginger_mono_yield_qntl_acr")),

        ("Ginger Poly",
         get_val(row, "agri_practices-Ginger_poly_acres"),
         get_val(row, "agri_practices-Ginger_poly_yield_qntl_acr")),

        ("Turmeric Mono",
         get_val(row, "agri_practices-Turmeric_mono_acres"),
         get_val(row, "agri_practices-Turmeric_mono_yield_qntl_acr")),

        ("Turmeric Poly",
         get_val(row, "agri_practices-Turmeric_poly_acres"),
         get_val(row, "agri_practices-Turmeric_poly_yield_qntl_acr")),

        ("Redgram Mono",
         get_val(row, "agri_practices-Redgram_mono_acres"),
         get_val(row, "agri_practices-Redgram_mono_yield_qntl_acr")),

        ("Redgram Poly",
         get_val(row, "agri_practices-Redgram_poly_acres"),
         get_val(row, "agri_practices-Redgram_poly_yield_qntl_acr")),

        ("Rajma Broadcasting",
         get_val(row, "agri_practices-Rajma_broadcast_acres"),
         get_val(row, "agri_practices-Rajma_broadcast_yield_qntl_acr")),

        ("Rajma Line Sowing",
         get_val(row, "agri_practices-Rajma_linesowing_acres"),
         get_val(row, "agri_practices-Rajma_linesowing_yield_qntl_acr")),
    ]

    crop_data = []

    for name, area, yield_val in crop_models_raw:
        area = pd.to_numeric(area, errors="coerce")
        yield_val = pd.to_numeric(yield_val, errors="coerce")

        area = float(area) if pd.notna(area) else 0
        yield_val = float(yield_val) if pd.notna(yield_val) else 0

        if area > 0 or yield_val > 0:
            crop_data.append((name, round(area, 2), round(yield_val, 2)))

    if crop_data:
        table_crop = doc.add_table(rows=len(crop_data) + 1, cols=3)
        table_crop.style = "Table Grid"

        table_crop.rows[0].cells[0].text = "Crop Model"
        table_crop.rows[0].cells[1].text = "Area (Acres)"
        table_crop.rows[0].cells[2].text = "Yield (Quintal/Acre)"

        for i, (name, area, yield_val) in enumerate(crop_data, start=1):
            table_crop.rows[i].cells[0].text = name
            table_crop.rows[i].cells[1].text = str(area)
            table_crop.rows[i].cells[2].text = str(yield_val)

    else:
        doc.add_paragraph("No crop model data available.")

    # -------- NATURAL FARMING PARAGRAPH --------
    doc.add_paragraph("")

    nf_hhs = pd.to_numeric(row.get("Total HH practicing NF", 0), errors="coerce")
    nf_area = pd.to_numeric(row.get("Total land under NF practice_acre", 0), errors="coerce")

    total_hhs = pd.to_numeric(row.get("Total HHs", 0), errors="coerce")
    total_land = pd.to_numeric(row.get("Total land in the village_acre", 0), errors="coerce")

    brc_available = str(row.get("brc available", "")).strip().lower()
    brc_name = row.get("Name of the BRC entrepreneur", "")

    villages_nf = pd.to_numeric(row.get("No of villages accessing NF inputs", 0), errors="coerce")
    farmers_nf = pd.to_numeric(row.get("No of farmers accessing NF inputs", 0), errors="coerce")
    brc_area = pd.to_numeric(row.get("Extent covered under BRC_acres", 0), errors="coerce")

    business_plan = str(row.get("Is business plan developed", "")).strip().lower()

    # safe conversion
    def safe_int(val):
        return int(val) if pd.notna(val) else 0

    nf_hhs = safe_int(nf_hhs)
    nf_area = safe_int(nf_area)
    total_hhs = safe_int(total_hhs)
    total_land = safe_int(total_land)
    villages_nf = safe_int(villages_nf)
    farmers_nf = safe_int(farmers_nf)
    brc_area = safe_int(brc_area)

    # -------- LOGIC --------
    if nf_hhs == 0 and nf_area == 0:

        nf_para = "There is no natural farming being practiced in this village."

    else:

        nf_para = f"""Natural farming is being practiced in the village, with {nf_hhs} households out of {total_hhs} households adopting natural farming over an extent of {nf_area} acres out of the total {total_land} acres in the village."""

    # 👉 ONLY IF BRC AVAILABLE
    if brc_available in ["yes", "y"]:
        nf_para += f""" A Bio-Resource Center (BRC) is available in the village, managed by {brc_name}, supporting {villages_nf} villages and {farmers_nf} farmers with natural farming inputs, covering an extent of {brc_area} acres."""

    # 👉 BUSINESS PLAN LOGIC (INSIDE BRC)
        if business_plan in ["yes", "y"]:
            nf_para += " The entrepreneur has developed a business plan for the BRC."
        else:
            nf_para += " The entrepreneur needs to develop a business plan to strengthen the BRC operations."

        # add to doc
        doc.add_paragraph(nf_para)

    # -------- NF ACTIVITY TABLE --------
    doc.add_heading("Natural Farming Activities", 2)

    nf_activity_raw = [
        ("PMDS", row.get("nf_activities-pmds_hhs", 0), row.get("nf_activities-pmds_acres", 0)),
        ("365DCC", row.get("nf_activities-_365dcc_hhs", 0), row.get("nf_activities-_365dcc_acres", 0)),
        ("RDS", row.get("nf_activities-rds_hhs", 0), row.get("nf_activities-rds_acres", 0)),
        ("5 Layer", row.get("nf_activities-_5layer_hhs", 0), row.get("nf_activities-_5layer_acres", 0)),
        ("Ghana Jeevamrutham", row.get("nf_activities-gja_hhs", 0), row.get("nf_activities-gja_acres", 0)),
        ("Dharajeeramrutham", row.get("nf_activities-dja_hhs", 0), row.get("nf_activities-dja_acres", 0)),
        ("Bheejamrutham", row.get("nf_activities-bheeejamruth_hhs", 0), row.get("nf_activities-bheeejamruth_acres", 0)),
        ("Concoctions", row.get("nf_activities-concoctions_hhs", 0), row.get("nf_activities-concoctions_acres", 0)),
        ("Mulching", row.get("nf_activities-mulching_hhs", 0), row.get("nf_activities-mulching_acres", 0)),
    ]

    nf_data = []

    for name, hhs, acres in nf_activity_raw:
        hhs = pd.to_numeric(hhs, errors="coerce")
        acres = pd.to_numeric(acres, errors="coerce")

        hhs = int(hhs) if pd.notna(hhs) else 0
        acres = int(acres) if pd.notna(acres) else 0

        # ✅ skip zero rows
        if hhs > 0 or acres > 0:
            nf_data.append((name, hhs, acres))

    # create table only if data exists
    if nf_data:
        table_nf = doc.add_table(rows=len(nf_data) + 1, cols=3)
        table_nf.style = "Table Grid"

        # header
        table_nf.rows[0].cells[0].text = "NF Activity"
        table_nf.rows[0].cells[1].text = "Farmers"
        table_nf.rows[0].cells[2].text = "Acres"

        # data rows
        for i, (name, hhs, acres) in enumerate(nf_data, start=1):
            table_nf.rows[i].cells[0].text = name
            table_nf.rows[i].cells[1].text = str(hhs)
            table_nf.rows[i].cells[2].text = str(acres)

    else:
        doc.add_paragraph("No natural farming activity available.")

    # -------- AGRI SERVICE & PROCESSING --------
    doc.add_paragraph("")
    doc.add_heading("Agri Service Center & Processing Facilities", 1)

    storage = str(row.get("infrastructure-storage_facility", "")).strip().lower()
    post_harvest = str(row.get("infrastructure-post_harvest_facility", "")).strip().lower()
    value_add = str(row.get("infrastructure-value_adding_facility", "")).strip().lower()

    asc = str(row.get("farmeasy_asc", "")).strip().lower()
    asc_plan = str(row.get("farmeasy_asc_businessplan", "")).strip().lower()

    # -------- LOGIC --------
    if asc in ["yes", "y"]:

        para_asc = "An Agri Service Center (ASC) is available in the village, supporting agricultural activities and farmer services."

        # facilities
        if storage == "yes":
            para_asc += " Storage facilities for agricultural produce are available."
        else:
            para_asc += " Storage facilities for agricultural produce are not available."

        if post_harvest == "yes":
            para_asc += " Post-harvest processing facilities are available."
        else:
            para_asc += " Post-harvest processing facilities are not available."

        if value_add == "yes":
            para_asc += " Value addition facilities are available."
        else:
            para_asc += " Value addition facilities are not available."

        # business plan
        if asc_plan in ["yes", "y"]:
            para_asc += " The Agri Service Center has a business plan in place."
        else:
            para_asc += " A business plan needs to be developed to strengthen the Agri Service Center."

    else:
        para_asc = "No Agri Service Center is available in the village, and there is a need to establish one to support farmers and improve agricultural services."

    doc.add_paragraph(para_asc)

    # -------- PROCESSING FACILITIES TABLE --------
    doc.add_paragraph("")
    doc.add_heading("Processing & Farm Mechanization Facilities", 1)

    facilities = [
        ("Power weeder", "farmeasy_Power_weeder_available", "farmeasy_Power_weeder_users"),
        ("Cycle weeder", "farmeasy_Cycle_weeder_available", "farmeasy_Cycle_weeder_users"),
        ("Plastic drums", "farmeasy_Plastic_drums_available", "farmeasy_Plastic_drums_users"),
        ("Cono weeder", "farmeasy_Cono_Weeder_available", "farmeasy_Cono_Weeder_users"),
        ("Sprayers", "farmeasy_Sprayers_available", "farmeasy_Sprayers_users"),
        ("Power sprayers", "farmeasy_Power_sprayers_available", "farmeasy_Power_sprayers_users"),
        ("Tarpaulin", "farmeasy_Taurplin_available", "farmeasy_Taurplin_users"),
        ("Graders", "farmeasy_Graders_available", "farmeasy_Graders_users"),
        ("Power tiller", "farmeasy_Power_tiller_available", "farmeasy_Power_tiller_users"),
        ("Tractor", "farmeasy_Tractor_available", "farmeasy_Tractor_users"),
        ("Coffee pulper", "farmeasy_Coffee_pulper_available", "farmeasy_Coffee_pulper_users"),
        ("Pepper thresher", "farmeasy_Pepper_thresher_available", "farmeasy_Pepper_thresher_users"),
        ("Multigrain thresher", "farmeasy_Multigrain_thresher_available", "farmeasy_Multigrain_thresher_users"),
        ("Turmeric polisher", "farmeasy_Turmeric_polisher_available", "farmeasy_Turmeric_polisher_users"),
        ("Turmeric boiler", "farmeasy_Turmeric_boiler_available", "farmeasy_Turmeric_boiler_users"),
        ("Micro dehullers", "processing_facilities-Millet_mixer", "processing_facilities-Millet_mixer_users"),
        ("Pulveriser", "processing_facilities-Pulveriser", "processing_facilities-Pulveriser_users"),
    ]

    # create table
    table_fac = doc.add_table(rows=len(facilities) + 1, cols=3)
    table_fac.style = "Table Grid"

    # header
    table_fac.rows[0].cells[0].text = "Facility"
    table_fac.rows[0].cells[1].text = "Units Available"
    table_fac.rows[0].cells[2].text = "Households Using"

    # rows
    for i, (name, col_avail, col_users) in enumerate(facilities, start=1):

        avail = pd.to_numeric(row.get(col_avail, 0), errors="coerce")
        users = pd.to_numeric(row.get(col_users, 0), errors="coerce")

        avail = int(avail) if pd.notna(avail) else 0
        users = int(users) if pd.notna(users) else 0

        table_fac.rows[i].cells[0].text = name
        table_fac.rows[i].cells[1].text = str(avail)
        table_fac.rows[i].cells[2].text = str(users)

        # -------- Livestock-Poultry-LR-SR-Fish --------
    doc.add_paragraph("")
    doc.add_heading("Poultry", 2)

    byp_hhs = pd.to_numeric(row.get("no of BYP HHs", 0), errors="coerce")
    birds = pd.to_numeric(row.get("Total Birds", 0), errors="coerce")
    mortality = pd.to_numeric(row.get("birds mortality", 0), errors="coerce")
    immunized = pd.to_numeric(row.get("Total birds immunized", 0), errors="coerce")

    service = str(row.get("poultry service provider", "")).strip().lower()

    def safe_int(val):
        return int(val) if pd.notna(val) else 0

    byp_hhs = safe_int(byp_hhs)
    birds = safe_int(birds)
    mortality = safe_int(mortality)
    immunized = safe_int(immunized)

    if byp_hhs == 0 and birds == 0:
        poultry_para = "Poultry activity is not significantly practiced in the village."
    else:
        poultry_para = f"""Backyard poultry is practiced by {byp_hhs} households in the village, with a total population of {birds} birds. During the last year, {mortality} birds were reported dead. Out of the total birds, {immunized} birds have been immunized."""

        if service in ["yes", "y"]:
            poultry_para += " Poultry service providers are available in the village."
        else:
            poultry_para += " There is a need to strengthen poultry services as no service provider is available."

    doc.add_paragraph(poultry_para)

    # -------- POULTRY TABLE --------
    poultry_data = [
        ("Households with Backyard Poultry", byp_hhs),
        ("Total Birds", birds),
        ("Bird Mortality (Last Year)", mortality),
        ("Birds Immunized", immunized),
        ("Service Provider Available", "Yes" if service in ["yes","y"] else "No"),
    ]

    table_poultry = doc.add_table(rows=len(poultry_data), cols=2)
    table_poultry.style = "Table Grid"

    for i, (name, val) in enumerate(poultry_data):
        table_poultry.rows[i].cells[0].text = name
        table_poultry.rows[i].cells[1].text = str(val)

    # -------- LARGE RUMINANTS --------
    doc.add_paragraph("")
    doc.add_heading("Large Ruminants", 2)

    total_hhs = pd.to_numeric(row.get("Total HHs", 0), errors="coerce")

    cows_hhs = pd.to_numeric(row.get("livestock-cows_hhs", 0), errors="coerce")
    bull_hhs = pd.to_numeric(row.get("livestock-bullocks_hhs", 0), errors="coerce")
    buff_hhs = pd.to_numeric(row.get("livestock-baffaloes_hhs", 0), errors="coerce")

    cows = pd.to_numeric(row.get("livestock-cows_nos", 0), errors="coerce")
    bulls = pd.to_numeric(row.get("livestock-bullocks_nos", 0), errors="coerce")
    buff = pd.to_numeric(row.get("livestock-baffaloes_nos", 0), errors="coerce")

    immunized = pd.to_numeric(row.get("Total animals immunized", 0), errors="coerce")
    mortality = pd.to_numeric(row.get("Mortality", 0), errors="coerce")

    sheds = pd.to_numeric(row.get("no of cattle sheds", 0), errors="coerce")
    renovated = pd.to_numeric(row.get("no of sheds renovated", 0), errors="coerce")

    def safe_int(val):
        return int(val) if pd.notna(val) else 0

    total_hhs = safe_int(total_hhs)
    cows_hhs, bull_hhs, buff_hhs = safe_int(cows_hhs), safe_int(bull_hhs), safe_int(buff_hhs)
    cows, bulls, buff = safe_int(cows), safe_int(bulls), safe_int(buff)
    immunized, mortality = safe_int(immunized), safe_int(mortality)
    sheds, renovated = safe_int(sheds), safe_int(renovated)

    total_animals = cows + bulls + buff
    total_livestock_hhs = cows_hhs + bull_hhs + buff_hhs

    # -------- LOGIC --------
    if total_animals == 0:
        lr_para = "Large ruminant rearing is not significantly practiced in the village."
    else:
        lr_para = f"""Large ruminants are reared by {total_livestock_hhs} households out of {total_hhs} households in the village, with a total of {total_animals} animals. During the last year, {mortality} animal deaths were reported."""

        # sheds
        if sheds > 0:
            lr_para += f". The village has {sheds} cattle sheds, of which {renovated} have been renovated"

        # immunization
        if immunized > 0:
            lr_para += f". A total of {immunized} animals have been immunized"
        else:
            lr_para += ". Animal immunization needs to be strengthened"

        lr_para += "."

    doc.add_paragraph(lr_para)  

    # -------- LARGE RUMINANTS TABLE --------
    lr_data = [
        ("Households with Cows", cows_hhs),
        ("Total Cows", cows),
        ("Households with Bullocks", bull_hhs),
        ("Total Bullocks", bulls),
        ("Households with Buffaloes", buff_hhs),
        ("Total Buffaloes", buff),
    ]

    table_lr = doc.add_table(rows=len(lr_data), cols=2)
    table_lr.style = "Table Grid"

    for i, (name, val) in enumerate(lr_data):
        table_lr.rows[i].cells[0].text = name
        table_lr.rows[i].cells[1].text = str(val)

    # -------- SMALL RUMINANTS --------
    doc.add_paragraph("")
    doc.add_heading("Small Ruminants", 2)

    total_hhs = pd.to_numeric(row.get("Total HHs", 0), errors="coerce")

    sr_hhs = pd.to_numeric(get_val(row, "livestock-sheep_goats_hhs"), errors="coerce")
    sr_animals = pd.to_numeric(get_val(row, "livestock-sheep_goats_nos"), errors="coerce")

    mortality = pd.to_numeric(get_val(row, "SR Mortality"), errors="coerce")
    immunized = pd.to_numeric(get_val(row, "SR immunized"), errors="coerce")

    sheds = pd.to_numeric(get_val(row, "no of sheep / goat sheds"), errors="coerce")
    elevated = pd.to_numeric(get_val(row, "no of elevated sheds"), errors="coerce")

    def safe_int(val):
        return int(val) if pd.notna(val) else 0

    total_hhs = safe_int(total_hhs)
    sr_hhs = safe_int(sr_hhs)
    sr_animals = safe_int(sr_animals)
    mortality = safe_int(mortality)
    immunized = safe_int(immunized)
    sheds = safe_int(sheds)
    elevated = safe_int(elevated)

    # -------- LOGIC --------
    if sr_animals == 0:
        sr_para = "Small ruminant rearing is not significantly practiced in the village."
    else:
        sr_para = f"""Small ruminants are reared by {sr_hhs} households out of {total_hhs} households in the village, with a total of {sr_animals} animals. During the last year, {mortality} animal deaths were reported"""

        if sheds > 0:
            sr_para += f". The village has {sheds} sheds, including {elevated} elevated sheds"

        if immunized > 0:
            sr_para += f". A total of {immunized} animals have been immunized"
        else:
            sr_para += ". Animal health services need strengthening, particularly immunization"

        sr_para += "."

    doc.add_paragraph(sr_para)

    # -------- SMALL RUMINANTS TABLE --------
    sr_data = [
        ("Households with Small Ruminants", sr_hhs),
        ("Total Sheep & Goats", sr_animals),
        ("Animal Mortality (Last Year)", mortality),
        ("Animals Immunized", immunized),
        ("Sheep/Goat Sheds", sheds),
        ("Elevated Sheds", elevated),
    ]

    table_sr = doc.add_table(rows=len(sr_data), cols=2)
    table_sr.style = "Table Grid"

    for i, (name, val) in enumerate(sr_data):
        table_sr.rows[i].cells[0].text = name
        table_sr.rows[i].cells[1].text = str(val)

        # -------- FISHERIES --------
    doc.add_paragraph("")
    doc.add_heading("Fisheries", 2)

    pond_hhs = pd.to_numeric(get_val(row, "HHs owning ponds"), errors="coerce")
    area = pd.to_numeric(get_val(row, "fisheries-area_ponds"), errors="coerce")
    water_spread = pd.to_numeric(get_val(row, "Total water spread are of the ponds_acr"), errors="coerce")

    individual = pd.to_numeric(get_val(row, "fisheries-individual_ponds"), errors="coerce")
    community = pd.to_numeric(get_val(row, "No of community ponds"), errors="coerce")

    farmponds = pd.to_numeric(get_val(row, "fisheries-eco_farmponds"), errors="coerce")
    total_ponds = pd.to_numeric(get_val(row, "No of ponds under fisheries"), errors="coerce")

    def safe_int(val):
        return int(val) if pd.notna(val) else 0

    pond_hhs = safe_int(pond_hhs)
    area = safe_int(area)
    water_spread = safe_int(water_spread)
    individual = safe_int(individual)
    community = safe_int(community)
    farmponds = safe_int(farmponds)
    total_ponds = safe_int(total_ponds)

    # -------- LOGIC --------
    if total_ponds == 0 and individual == 0:
        fish_para = "Fisheries activities are not significantly practiced in the village."
    else:
        fish_para = f"""Fisheries activities are practiced in the village, with {pond_hhs} households owning ponds. The total area under ponds is {area} acres, with a water spread area of {water_spread} acres. The village has {individual} individual ponds and {community} community ponds."""

        if farmponds > 0:
            fish_para += f" Additionally, {farmponds} eco-farm ponds are available"

        if total_ponds > 0:
            fish_para += f", with a total of {total_ponds} ponds under fisheries"

        fish_para += "."

    doc.add_paragraph(fish_para)

    # -------- FISHERIES TABLE --------
    fish_data = [
        ("Households owning ponds", pond_hhs),
        ("Area under ponds (acres)", area),
        ("Water spread area (acres)", water_spread),
        ("Individual ponds", individual),
        ("Community ponds", community),
        ("Eco farm ponds", farmponds),
        ("Total ponds under fisheries", total_ponds),
    ]

    table_fish = doc.add_table(rows=len(fish_data), cols=2)
    table_fish.style = "Table Grid"

    for i, (name, val) in enumerate(fish_data):
        table_fish.rows[i].cells[0].text = name
        table_fish.rows[i].cells[1].text = str(val)

# -------- WATER RESOURCE DEVELOPMENT --------
    doc.add_paragraph("")
    doc.add_heading("Water Resource Development", 1)

    def get_water_val(row, col):
        cols = list(row.index)
        if col in cols:
            return row[col]
        for c in cols:
            if col.replace("-", "").replace("_", "") == c.replace("-", "").replace("_", ""):
                return row[c]
        return 0

    checkdam = str(get_water_val(row, "checkdam-checkdam_in_village")).strip().lower()
    repair = str(get_water_val(row, "checkdam-checkdam_repairs")).strip().lower()

    year = str(get_water_val(row, "checkdam-Checkdam_year"))
    before = pd.to_numeric(get_water_val(row, "checkdam-actual_coverage_before_repair"), errors="coerce")
    after = pd.to_numeric(get_water_val(row, "checkdam-expected_coverage_after_repair"), errors="coerce")

    before = int(before) if pd.notna(before) else 0
    after = int(after) if pd.notna(after) else 0

    if checkdam in ["yes", "y"]:
        para_water = "A check dam is available in the village"

        if year and year != "0":
            para_water += ", constructed in " + year[:4]

        if repair in ["yes", "y"]:
            para_water += ". The check dam requires repairs"

        if before > 0 or after > 0:
            para_water += ", irrigation coverage is " + str(before) + " acres and may increase to " + str(after) + " acres"

        para_water += ". This structure supports water availability and agriculture."
    else:
        para_water = "No check dam is available in the village. There is a need for water harvesting structures."

    doc.add_paragraph(para_water)

    # -------- SHG & FPO TABLE --------
    doc.add_paragraph("")
    doc.add_heading("SHGs and Farmer Producer Organizations", 1)

    shg = pd.to_numeric(get_val(row, "SHG_nos"), errors="coerce")
    fpo_groups = pd.to_numeric(get_val(row, "fpo-farmer_groups"), errors="coerce")
    fpo_member_flag = str(get_val(row, "fpo-fpo_fpc_membership")).strip()
    fpo_name = str(get_val(row, "fpo-fpo_fpc_name"))
    fpo_members = pd.to_numeric(get_val(row, "fpo-fpo_fpc_members"), errors="coerce")

    # safe conversion
    def safe_int(val):
        return int(val) if pd.notna(val) else 0

    shg = safe_int(shg)
    fpo_groups = safe_int(fpo_groups)
    fpo_members = safe_int(fpo_members)

    # table data (no filtering)
    shg_data = [
        ("No. of SHGs", shg),
        ("No. of Farmer Producer Groups", fpo_groups),
        ("FPO/FPC Membership Available", fpo_member_flag),
        ("FPO/FPC Name", fpo_name),
        ("No. of FPO Members", fpo_members),
    ]

    table_shg = doc.add_table(rows=len(shg_data), cols=2)
    table_shg.style = "Table Grid"

    for i, (name, val) in enumerate(shg_data):
        table_shg.rows[i].cells[0].text = str(name)
        table_shg.rows[i].cells[1].text = str(val)

        # -------- OTHER WORKS --------
    doc.add_paragraph("")
    doc.add_heading("Other Development Works Identified", 1)

    other = str(get_val(row, "other_info")).strip()

    if other and other != "0":
        para_other = "The following additional development works have been identified in the village: " + other
    else:
        para_other = "No additional development works were identified in the village."

    doc.add_paragraph(para_other)

    # -------- ENUMERATOR DETAILS --------
    doc.add_paragraph("")
    doc.add_heading("Data Collection Details", 1)

    enum_name = str(get_val(row, "enumerator-Enumerator_name"))
    enum_mobile = str(get_val(row, "enumerator-Enumerator_mobile"))

    enum_data = [
        ("Enumerator Name", enum_name),
        ("Contact Number", enum_mobile),
    ]

    table_enum = doc.add_table(rows=len(enum_data), cols=2)
    table_enum.style = "Table Grid"

    for i, (name, val) in enumerate(enum_data):
        table_enum.rows[i].cells[0].text = name
        table_enum.rows[i].cells[1].text = val

   # -------- PROPOSED DEVELOPMENT INTERVENTIONS --------
    doc.add_paragraph("")
    doc.add_heading("Proposed Development Interventions", 1)

    interventions = []
    total_cost = 0

    for _, row_b in df_budget.iterrows():

        source_col = row_b["Source Column"]
        work_name = row_b["Name of the work"]
        unit_cost = pd.to_numeric(row_b["Unit Cost (Rs)"], errors="coerce")

        # check column exists in village plan
        if source_col in df_p.columns:

            qty = pd.to_numeric(df_p[source_col], errors="coerce").sum()

            if qty > 0:
                cost = qty * unit_cost
                total_cost += cost

                interventions.append(work_name)

    # -------- CREATE SMART TEXT --------
    if len(interventions) > 0:

        # join works nicely
        if len(interventions) == 1:
            works_text = interventions[0]
        else:
            works_text = ", ".join(interventions[:-1]) + " and " + interventions[-1]

        # convert cost to lakhs
        total_cost_lakh = total_cost / 100000

        para_plan = f"""Based on the village assessment and identified gaps, a comprehensive set of development interventions has been proposed to strengthen livelihoods, agriculture, and allied sectors. Key interventions include {works_text}. The total estimated investment required for these interventions is approximately ₹{total_cost_lakh:.2f} lakhs. These interventions are expected to enhance productivity, improve infrastructure, and create sustainable income opportunities for households."""

    else:
        para_plan = "No specific development interventions have been proposed for the village based on the current assessment."

    doc.add_paragraph(para_plan)

    # -------- DETAILED INTERVENTIONS TABLE --------
    doc.add_paragraph("")
    doc.add_heading("Detailed Proposed Works", 2)

    table_data = []

    for _, row_b in df_budget.iterrows():

        source_col = row_b["Source Column"]
        work_name = row_b["Name of the work"]
        unit_cost = pd.to_numeric(row_b["Unit Cost (Rs)"], errors="coerce")

        if source_col in df_p.columns:

            qty = pd.to_numeric(df_p[source_col], errors="coerce").sum()

            if qty > 0:
                total = qty * unit_cost

                table_data.append([
                    work_name,
                    int(qty),
                    int(unit_cost),
                    int(total)
                ])

    # create table
    table = doc.add_table(rows=len(table_data)+1, cols=4)
    table.style = "Table Grid"

    # headers
    headers = ["Work", "Units", "Unit Cost (₹)", "Total Cost (₹)"]

    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h

    # fill rows
    for i, row_data in enumerate(table_data):
        for j, val in enumerate(row_data):
            table.rows[i+1].cells[j].text = str(val)


    # -------- OTHER SECTIONS --------

    # -------- VILLAGE PHOTOS --------
    doc.add_paragraph("")
    doc.add_heading("Village Photos", 1)

    import os

    image_found = False

    for i in range(1, 6):  # supports up to 5 images (you can increase)

        image_path = f"images/{village}_{i}.jpg"

        if os.path.exists(image_path):
            doc.add_picture(image_path, width=Inches(5.5))
            doc.add_paragraph("")  # spacing
            image_found = True

    # fallback
    if not image_found:
        doc.add_paragraph("No images available for this village.")

    # SAVE
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


# ---------------- BUTTON ----------------
if st.button("Generate Report"):

    report = generate_report(df_v, df_p, selected_village)
    file = create_doc(report, df_v, selected_village)

    st.download_button(
        "Download Report",
        data=file,
        file_name=f"{selected_village}_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
